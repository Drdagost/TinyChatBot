from docx import Document
from docx.enum.text import WD_BREAK

from tinychatbot.io_utils import DocumentExtractor


def make_multipage_docx(path: str):
    doc = Document()
    # Page 1 content
    doc.add_paragraph('PageOne-Intro')
    doc.add_paragraph('PageOne-Body')
    # Insert explicit page break
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)

    # Page 2 content
    doc.add_paragraph('PageTwo-Intro')
    doc.add_paragraph('PageTwo-Body')

    # Add a table on page 2
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'p2r0c0'
    table.cell(0, 1).text = 'p2r0c1'
    table.cell(1, 0).text = 'p2r1c0'
    table.cell(1, 1).text = 'p2r1c1'

    doc.save(path)


def test_docx_multipage_extraction(tmp_path):
    p = tmp_path / 'multipage.docx'
    make_multipage_docx(str(p))

    extractor = DocumentExtractor(enable_ocr=False)
    text = extractor.extract(str(p))

    # Assert text from both pages is present
    assert 'PageOne-Intro' in text
    assert 'PageOne-Body' in text
    assert 'PageTwo-Intro' in text
    assert 'PageTwo-Body' in text
    # Assert table content appears
    assert 'p2r0c0' in text and 'p2r1c1' in text
