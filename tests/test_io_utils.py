from docx import Document

from tinychatbot.io_utils import DocumentExtractor


def make_docx(path: str):
    doc = Document()
    # header
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "HEADER_TEXT_123"

    # footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "FOOTER_TEXT_456"

    # paragraph
    doc.add_paragraph("Hello world paragraph")

    # table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'r0c0'
    table.cell(0, 1).text = 'r0c1'
    table.cell(1, 0).text = 'r1c0'
    table.cell(1, 1).text = 'r1c1'

    doc.save(path)


def test_docx_extraction_includes_header_footer_and_table(tmp_path):
    p = tmp_path / 'test_doc.docx'
    make_docx(str(p))

    extractor = DocumentExtractor(enable_ocr=False)
    text = extractor.extract(str(p))

    assert 'HEADER_TEXT_123' in text
    assert 'FOOTER_TEXT_456' in text
    assert 'Hello world paragraph' in text
    assert 'r0c0' in text and 'r1c1' in text
