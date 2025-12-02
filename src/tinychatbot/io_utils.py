import os
import shutil
from typing import Callable, Dict, List, Optional

from loguru import logger


def check_native_binaries() -> Dict[str, bool]:
    """Return flags indicating presence of external binaries used for OCR and PDF->image conversion.

    Also prints a clear warning if either is missing.
    """
    bins = {
        "tesseract": shutil.which("tesseract") is not None,
        "pdftoppm": shutil.which("pdftoppm") is not None,
    }

    if not bins["tesseract"] or not bins["pdftoppm"]:
        lines = ["\n[WARNING] Native binaries missing for OCR/text extraction:"]
        if not bins["tesseract"]:
            lines.append(" - tesseract (required by pytesseract) not found in PATH.")
        if not bins["pdftoppm"]:
            lines.append(
                " - pdftoppm (part of poppler) not found in PATH; required by pdf2image."
            )
        lines.append(
            "\nOCR and image-based PDF text extraction will be skipped unless you install the missing binaries."
        )
        lines.append("See the README for Windows installation instructions.")
        logger.warning("\n".join(lines))

    return bins


def _extract_docx_text_simple(path: str) -> str:
    try:
        from docx import Document
    except Exception:
        return ""
    try:
        doc = Document(path)
    except Exception:
        return ""
    parts = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(parts)


def _extract_pdf_text_simple(path: str) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        text = "\n".join([p.extract_text() or "" for p in reader.pages])
        return text
    except Exception:
        return ""


class DocumentExtractor:
    """Pluggable document text extractor.

    Basic usage:
      extractor = DocumentExtractor(enable_ocr=None)
      text = extractor.extract(path)
      docs = extractor.load_folder(folder_path)

    enable_ocr: None = auto-detect via `check_native_binaries()`; True/False to force.
    """

    def __init__(self, enable_ocr: Optional[bool] = None):
        self.native_bins = check_native_binaries()
        if enable_ocr is None:
            self.enable_ocr = bool(
                self.native_bins.get("tesseract") and self.native_bins.get("pdftoppm")
            )
        else:
            self.enable_ocr = bool(enable_ocr)
        # Handler registry keyed by extension
        self.handlers: Dict[str, Callable[[str], str]] = {
            ".pdf": self._handle_pdf,
            ".docx": self._handle_docx,
            ".txt": self._handle_text,
            ".md": self._handle_text,
        }

    def register_handler(self, ext: str, fn: Callable[[str], str]):
        """Register a custom handler for file extension (ext should include leading dot)."""
        self.handlers[ext.lower()] = fn

    def extract(self, path: str) -> str:
        """Extract text from a single file path using the registered handlers and sensible fallbacks."""
        _, ext = os.path.splitext(path)
        ext = ext.lower()
        handler = self.handlers.get(ext)
        if handler:
            try:
                return handler(path) or ""
            except Exception:
                return ""

        # Generic fallback: best-effort open as text
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""

    def load_folder(self, folder_path: str) -> List[Dict[str, str]]:
        """Walk a folder and return list of {'path': path, 'text': text} for readable documents."""
        docs = []
        for root, _, files in os.walk(folder_path):
            for fname in files:
                path = os.path.join(root, fname)
                text = self.extract(path)
                docs.append({"path": path, "text": text})
        return docs

    # --- Handlers ---
    def _handle_text(self, path: str) -> str:
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception:
            return ""

    def _handle_docx(self, path: str) -> str:
        # richer extraction: collect text from all package parts to include textboxes/shapes
        try:
            from docx import Document
        except Exception:
            return _extract_docx_text_simple(path)

        try:
            doc = Document(path)
        except Exception:
            return ""

        parts = []

        # Iterate over all parts in the package (document, headers, footers, footnotes, endnotes, etc.)
        try:
            package = doc.part.package
            for part in package.parts:
                try:
                    for node in part.element.iter():
                        # collect text nodes (w:t) used by WordprocessingML and a:t in drawing/textbox
                        tag = node.tag
                        if tag is None:
                            continue
                        localname = tag.split("}")[-1]
                        if localname == "t" and node.text:
                            parts.append(node.text)
                except Exception:
                    continue
        except Exception:
            # Fallback to simple paragraph/table/header extraction
            for p in doc.paragraphs:
                if p.text:
                    parts.append(p.text)
            try:
                for section in doc.sections:
                    for p in section.header.paragraphs:
                        if p.text:
                            parts.append(p.text)
                    for p in section.footer.paragraphs:
                        if p.text:
                            parts.append(p.text)
            except Exception:
                pass
            try:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells if cell.text]
                        if row_text:
                            parts.append(" | ".join(row_text))
            except Exception:
                pass

        # Inline shapes alt text fallback (if present)
        try:
            if hasattr(doc, "inline_shapes"):
                for shape in doc.inline_shapes:
                    try:
                        if hasattr(shape, "_inline") and hasattr(
                            shape._inline, "docPr"
                        ):
                            docpr = shape._inline.docPr
                            alt = getattr(docpr, "descr", None) or getattr(
                                docpr, "title", None
                            )
                            if alt:
                                parts.append(alt)
                    except Exception:
                        continue
        except Exception:
            pass

        return "\n".join([p for p in parts if p])

    def _handle_pdf(self, path: str) -> str:
        text = _extract_pdf_text_simple(path)
        if text.strip():
            return text
        if not self.enable_ocr:
            return text

        # Try OCR path
        try:
            import pytesseract  # type: ignore
            from pdf2image import convert_from_path  # type: ignore

            images = convert_from_path(path)
            texts = []
            for img in images:
                try:
                    texts.append(pytesseract.image_to_string(img))
                except Exception:
                    texts.append("")
            return "\n".join(texts)
        except Exception:
            return text


# Backwards-compatible helpers
def extract_docx_text(path: str) -> str:
    return DocumentExtractor().extract(path)


def extract_pdf_text_with_ocr(path: str, enable_ocr: bool = True) -> str:
    return DocumentExtractor(enable_ocr=enable_ocr).extract(path)
